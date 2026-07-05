"""Renders agent-authored content into real, downloadable .docx/.pdf files."""

from __future__ import annotations

import io
from dataclasses import dataclass

from docx import Document as DocxDocument
from fpdf import FPDF
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.base import new_uuid
from app.db.models.document import Document
from app.repositories.document_repository import DocumentRepository
from app.services.storage_service import get_storage_backend


@dataclass
class DocumentSection:
    heading: str
    body: str


def _safe_pdf_text(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def _render_docx(title: str, sections: list[DocumentSection]) -> bytes:
    doc = DocxDocument()
    doc.add_heading(title, level=0)
    for section in sections:
        doc.add_heading(section.heading, level=1)
        for paragraph in section.body.split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _write_block(pdf: FPDF, h: int, text: str) -> None:
    # multi_cell defaults to leaving the cursor at the right edge of the last line
    # (new_x=XPos.RIGHT), which starves the *next* multi_cell of horizontal space and
    # raises FPDFException. Force the cursor back to the left margin on a new line.
    pdf.multi_cell(0, h, text, new_x="LMARGIN", new_y="NEXT")


def _render_pdf(title: str, sections: list[DocumentSection]) -> bytes:
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 18)
    _write_block(pdf, 10, _safe_pdf_text(title))
    pdf.ln(2)

    for section in sections:
        pdf.set_font("Helvetica", "B", 13)
        _write_block(pdf, 8, _safe_pdf_text(section.heading))
        pdf.set_font("Helvetica", "", 11)
        _write_block(pdf, 6, _safe_pdf_text(section.body))
        pdf.ln(3)

    return bytes(pdf.output())


class DocumentService:
    def __init__(self, session: AsyncSession):
        self.session = session
        self.repo = DocumentRepository(session)
        self.storage = get_storage_backend()

    async def create_document(
        self,
        user_id: str,
        title: str,
        doc_type: str,
        file_format: str,
        sections: list[DocumentSection],
    ) -> Document:
        if file_format == "pdf":
            content = _render_pdf(title, sections)
        else:
            file_format = "docx"
            content = _render_docx(title, sections)

        filename = f"{new_uuid()}.{file_format}"
        relative_path = f"{user_id}/{filename}"
        full_path = self.storage.save(relative_path, content)

        document = Document(
            user_id=user_id,
            title=title,
            doc_type=doc_type,
            file_format=file_format,
            file_path=full_path,
            status="ready",
        )
        await self.repo.add(document)
        return document
